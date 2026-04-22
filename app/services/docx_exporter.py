from __future__ import annotations

import io
from datetime import date
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _MPL_AVAILABLE = True
except ImportError:
    _MPL_AVAILABLE = False


class DocxExportError(RuntimeError):
    """Raised when Word document generation fails."""


def report_to_docx(report: dict[str, Any]) -> bytes:
    """Generate a comprehensive Word document from a report dict."""
    if not _DOCX_AVAILABLE:
        raise DocxExportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    _set_page_margins(doc)

    protein = report["protein"]
    variant = report["variant"]
    accession = protein.accession
    var_text = variant.short_notation
    gene = protein.gene_symbol or protein.entry_name or accession

    # ── Title Page ─────────────────────────────────────────────────────
    h = doc.add_heading("Protein Mutation Analysis Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"{gene}  ·  {var_text}", bold=True, size=16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, protein.protein_name or "", size=11, color=(80, 80, 80))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, f"Generated: {date.today().isoformat()}  ·  UniProt: {accession}", size=9, color=(120, 120, 120))

    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    _two_col_table(doc, [
        ("Accession", accession),
        ("Gene symbol", gene),
        ("Protein name", protein.protein_name or "—"),
        ("Organism", protein.organism_name or "—"),
        ("Sequence length", f"{protein.length} aa"),
        ("Variant submitted", variant.original_text),
        ("Variant normalized", var_text),
        ("HGVS protein", variant.hgvs_protein or "—"),
        ("Residue change", f"{variant.wild_type} → {variant.mutant} at position {variant.position}"),
    ])

    if report.get("summary_cards"):
        doc.add_heading("Key Findings", level=2)
        _two_col_table(doc, [(c["label"], str(c["value"])) for c in report["summary_cards"]])

    # ── Stability quick-verdict inline ─────────────────────────────────
    sp = report.get("stability_panel")
    if sp:
        verdict_p = doc.add_paragraph()
        _run(verdict_p, "Stability verdict: ", bold=True)
        _run(verdict_p, sp.get("summary", ""), size=10)

    # ── 2. Physicochemical Metrics ─────────────────────────────────────
    doc.add_heading("2. Physicochemical Metrics (WT vs Mutant)", level=1)

    if report.get("metric_deltas"):
        cols = ["Metric", "Wild-type", "Mutant", "Δ (Mutant − WT)"]
        rows = []
        for m in report["metric_deltas"]:
            wt  = f"{m.wild_type:.4f}" if m.wild_type  is not None else "N/A"
            mt  = f"{m.mutant:.4f}"    if m.mutant      is not None else "N/A"
            dlt = f"{m.delta:+.4f}"   if m.delta       is not None else "N/A"
            rows.append([m.label, wt, mt, dlt])
        _data_table(doc, cols, rows)

        if _MPL_AVAILABLE:
            img = _metric_delta_chart(report["metric_deltas"])
            if img:
                doc.add_paragraph("Physicochemical delta chart:")
                doc.add_picture(io.BytesIO(img), width=Inches(5.5))

    # ── 3. Residue Context ─────────────────────────────────────────────
    doc.add_heading("3. Residue Context", level=1)
    ctx = report.get("context") or {}
    if ctx:
        doc.add_paragraph(
            f"Sequence window: residues {ctx.get('window_start','?')} – {ctx.get('window_end','?')}"
        )
        p_wt = doc.add_paragraph()
        _run(p_wt, "Wild-type:  ", bold=True)
        _run(p_wt, ctx.get("wild_window", ""), mono=True)
        p_mt = doc.add_paragraph()
        _run(p_mt, "Mutant:     ", bold=True)
        _run(p_mt, ctx.get("mutant_window", ""), mono=True)

        wtp = ctx.get("wild_type_properties") or {}
        mtp = ctx.get("mutant_properties") or {}
        if wtp or mtp:
            doc.add_heading("Residue Properties at Variant Site", level=2)
            _two_col_table(doc, [
                ("WT amino acid", variant.wild_type),
                ("WT hydropathy",   str(wtp.get("hydropathy", "—"))),
                ("WT charge",       wtp.get("charge_class", "—")),
                ("WT polarity",     wtp.get("polarity_class", "—")),
                ("WT aromatic",     "Yes" if wtp.get("aromatic") else "No"),
                ("Mutant amino acid", variant.mutant),
                ("Mutant hydropathy", str(mtp.get("hydropathy", "—"))),
                ("Mutant charge",     mtp.get("charge_class", "—")),
                ("Mutant polarity",   mtp.get("polarity_class", "—")),
                ("Mutant aromatic",   "Yes" if mtp.get("aromatic") else "No"),
            ])

    # ── 4. Mutation Stability Prediction ───────────────────────────────
    if sp:
        doc.add_heading("4. Mutation Stability Prediction", level=1)
        _two_col_table(doc, [
            ("ΔΔG estimate (kcal/mol)", f"{sp['ddg_estimate']:+.2f}"),
            ("Classification",          sp.get("ddg_label", "—")),
            ("BLOSUM62 score",          str(sp.get("blosum62_score", "—"))),
            ("BLOSUM62 interpretation", sp.get("blosum62_interpretation", "—")),
            ("Hydrophobicity Δ (KD scale)", f"{sp.get('hydro_delta', 0):+.2f}"),
            ("Volume Δ (Å³)",           f"{sp.get('vol_delta', 0):+.1f}"),
            ("SIFT-proxy tolerance",    f"{sp.get('sift_proxy', '—')} ({sp.get('sift_tolerance', '—')})"),
        ])
        doc.add_paragraph(sp.get("summary", ""))

        if sp.get("property_changes"):
            doc.add_heading("Physicochemical Property Changes", level=2)
            _data_table(doc,
                ["Property", "From", "To", "Impact"],
                [[c["property"], c["from"], c["to"], c["impact"]] for c in sp["property_changes"]])

        disc = sp.get("disclaimer", "")
        if disc:
            p = doc.add_paragraph(disc)
            if p.runs:
                p.runs[0].italic = True

    # ── 5. Secondary Structure (PSIPRED) ───────────────────────────────
    ss = report.get("secondary_structure_panel")
    if ss:
        doc.add_heading("5. Secondary Structure Prediction (PSIPRED)", level=1)
        vr = ss.get("variant_row") or {}
        doc.add_paragraph(
            f"Variant site {variant.wild_type}{variant.position}{variant.mutant} is predicted as "
            f"{vr.get('state', '—')} with {vr.get('confidence_pct', '—')}% confidence."
        )
        p_prob = doc.add_paragraph()
        for card in ss.get("probability_cards", []):
            _run(p_prob, f"  {card['label']}: {card['value']}%  ", bold=False, size=9)

        doc.add_heading("State Distribution", level=2)
        _data_table(doc,
            ["State", "Count", "Percent (%)"],
            [[c["label"], str(c["count"]), f"{c['percent']}%"] for c in ss.get("state_cards", [])])

        if _MPL_AVAILABLE:
            img = _ss_pie_chart(ss.get("state_cards", []))
            if img:
                doc.add_picture(io.BytesIO(img), width=Inches(3.2))

        # Per-residue table (first 60 residues to keep doc manageable)
        rows_ss = ss.get("rows", [])
        if rows_ss:
            doc.add_heading("Per-residue Predictions (first 60 residues)", level=2)
            _data_table(doc,
                ["Position", "AA", "State", "Confidence (%)"],
                [[str(r.get("position","—")), r.get("aa","—"), r.get("state","—"), str(r.get("confidence_pct","—"))]
                 for r in rows_ss[:60]])
    else:
        doc.add_heading("5. Secondary Structure Prediction", level=1)
        doc.add_paragraph("PSIPRED predictor not configured or no output available. "
                          "Install PSIPRED and set PSIPRED_COMMAND_TEMPLATE in .env.predictors.")

    # ── 6. Solvent Accessibility (DMVFL-RSA) ───────────────────────────
    acc_panel = report.get("accessibility_panel")
    if acc_panel:
        doc.add_heading("6. Solvent Accessibility (DMVFL-RSA)", level=1)
        vr = acc_panel.get("variant_row") or {}
        _two_col_table(doc, [
            ("Mean RSA (whole protein)",    str(acc_panel.get("mean_rsa", "—"))),
            ("Mean ASA (Å²)",               str(acc_panel.get("mean_asa", "—"))),
            ("Exposed residues (RSA > 0.2)", str(acc_panel.get("exposed_count", "—"))),
            ("Buried residues (RSA ≤ 0.2)", str(acc_panel.get("buried_count", "—"))),
            (f"Variant RSA (pos. {variant.position})", str(vr.get("RSA", "—"))),
            (f"Variant ASA (Å²)",            str(vr.get("ASA", "—"))),
            ("Variant exposure class",      str(vr.get("Exposure", "—"))),
        ])
    else:
        doc.add_heading("6. Solvent Accessibility", level=1)
        doc.add_paragraph("DMVFL-RSA predictor not configured or no output available.")

    # ── 7. Integrated Structural Verdict ───────────────────────────────
    sv = report.get("structural_verdict")
    if sv:
        doc.add_heading("7. Integrated Structural Verdict", level=1)
        doc.add_paragraph(sv.get("verdict_sentence", ""))
        if sv.get("domain_context"):
            doc.add_paragraph(f"Domain context: {sv['domain_context']}")
        parts = sv.get("parts", [])
        if parts:
            p = doc.add_paragraph()
            _run(p, "Evidence pillars: ", bold=True)
            _run(p, "  ·  ".join(parts))

    ap = report.get("amphipathic_panel")
    if ap:
        doc.add_heading("Amphipathic Helix", level=2)
        _two_col_table(doc, [
            ("Segment",      f"{ap.get('segment_start','—')}–{ap.get('segment_end','—')}"),
            ("μH (moment)",  str(ap.get("mu_h", "—"))),
            ("Class",        ap.get("class_label", "—")),
            ("Summary",      ap.get("summary", "—")),
        ])

    # ── 8. AlphaFold2 Structural Confidence ────────────────────────────
    af = report.get("alphafold_panel")
    if af:
        doc.add_heading("8. AlphaFold2 Structural Confidence", level=1)
        doc.add_paragraph(
            f"Entry: {af.get('entry_id','—')}  ·  Model date: {af.get('model_date','—')}  ·  "
            f"Coverage: residues {af.get('uniprot_start','—')}–{af.get('uniprot_end','—')}"
        )
        if af.get("plddt_available"):
            vr = af.get("variant_row") or {}
            _two_col_table(doc, [
                ("Mean pLDDT (whole model)",          f"{af.get('mean_plddt','—')} — {af.get('mean_plddt_label','—')}"),
                (f"Variant pLDDT (pos. {variant.position})", f"{vr.get('plddt','—')} — {vr.get('plddt_label','—')}"),
            ])

            conf_cards = af.get("confidence_cards", [])
            if conf_cards:
                doc.add_heading("pLDDT Confidence Distribution", level=2)
                _data_table(doc,
                    ["Confidence band", "Count", "Percent (%)"],
                    [[c["label"], str(c["count"]), f"{c['percent']}%"] for c in conf_cards])

        am = af.get("alphamissense")
        if am:
            doc.add_heading("AlphaMissense Pathogenicity Score", level=2)
            _two_col_table(doc, [
                ("Score",       str(am.get("score", "—"))),
                ("Class",       am.get("label", "—")),
                ("Interpretation", am.get("interpretation", "—")),
            ])
    else:
        doc.add_heading("8. AlphaFold2 Structural Confidence", level=1)
        doc.add_paragraph("No AlphaFold entry found for this protein.")

    # ── 9. Experimental PDB Structures ─────────────────────────────────
    pdb_panel = report.get("pdb_panel")
    if pdb_panel:
        doc.add_heading("9. Experimental PDB Structures", level=1)
        doc.add_paragraph(
            f"{pdb_panel.get('total_count',0)} PDB entries found for {accession}. "
            f"{pdb_panel.get('variant_covered_count',0)} cover position {variant.position}."
        )
        best = pdb_panel.get("best_entry")
        if best:
            doc.add_heading("Best Structure at Variant Site", level=2)
            _two_col_table(doc, [
                ("PDB ID",      best.get("pdb_id", "—")),
                ("Method",      best.get("method", "—")),
                ("Resolution",  str(best.get("resolution", "—"))),
                ("Chain",       best.get("chain", "—")),
                ("Residue range", f"{best.get('residue_start','—')}–{best.get('residue_end','—')}"),
            ])
        entries = (pdb_panel.get("entries") or [])[:20]
        if entries:
            doc.add_heading("All PDB Entries (top 20)", level=2)
            _data_table(doc,
                ["PDB ID", "Method", "Resolution", "Chain", "Range", "Covers Variant"],
                [[e.get("pdb_id","—"), e.get("method","—"), str(e.get("resolution","—")),
                  e.get("chain","—"), f"{e.get('residue_start','—')}–{e.get('residue_end','—')}",
                  "Yes" if e.get("covers_variant") else "No"] for e in entries])
    else:
        doc.add_heading("9. Experimental PDB Structures", level=1)
        doc.add_paragraph("No experimental PDB cross-references found for this protein.")

    # ── 10. Primate Conservation ────────────────────────────────────────
    cons = report.get("conservation_panel")
    if cons:
        doc.add_heading("10. Primate Conservation", level=1)
        _two_col_table(doc, [
            ("Conservation at position",   f"{cons.get('conservation_pct','—')}%"),
            ("Human residue",              cons.get("human_residue", "—")),
            ("Dominant residue",           cons.get("dominant_residue", "—")),
            ("Conserved orthologues",      f"{cons.get('conserved_count','—')}/{cons.get('comparable_count','—')}"),
            ("Orthologues matching mutant", str(cons.get("mutant_match_count", "—"))),
        ])
        rows_c = cons.get("species_rows", [])
        if rows_c:
            doc.add_heading("Species Table", level=2)
            _data_table(doc,
                ["Species", f"Residue at pos. {cons.get('variant_position','?')}", "Match class"],
                [[r["species_label"], r["site_residue"], r["match_class"]] for r in rows_c])

    # ── 11. Genomic Location ────────────────────────────────────────────
    chr_panel = report.get("chromosome_location_panel")
    if chr_panel:
        doc.add_heading("11. Genomic Location", level=1)
        _two_col_table(doc, [
            ("Chromosome",   f"chr{chr_panel.get('chromosome','—')}"),
            ("Gene",         chr_panel.get("gene_symbol", "—")),
            ("Gene span",    chr_panel.get("gene_span_label", "—")),
            ("Strand",       chr_panel.get("strand_label", "—")),
            ("Transcript",   chr_panel.get("transcript_id", "—")),
            ("Variant span", chr_panel.get("variant_span_label", "—")),
        ])

    # ── 12. Human Protein Atlas ─────────────────────────────────────────
    pa = report.get("protein_atlas_panel")
    if pa:
        doc.add_heading("12. Human Protein Atlas", level=1)
        doc.add_paragraph(
            f"Source: Human Protein Atlas (HPA) — {pa.get('hpa_url','')}  ·  "
            f"Ensembl: {pa.get('ensembl_id','—')}  ·  Evidence: {pa.get('evidence','—')}"
        )
        if pa.get("gene_description"):
            doc.add_paragraph(pa["gene_description"])

        if pa.get("protein_class"):
            doc.add_heading("Protein Class", level=2)
            doc.add_paragraph(" · ".join(pa["protein_class"]))

        if pa.get("biological_process"):
            doc.add_heading("Biological Process", level=2)
            doc.add_paragraph(" · ".join(pa["biological_process"]))

        if pa.get("molecular_function"):
            doc.add_heading("Molecular Function", level=2)
            doc.add_paragraph(" · ".join(pa["molecular_function"]))

        if pa.get("subcellular_locations"):
            doc.add_heading("Subcellular Localisation", level=2)
            _data_table(doc,
                ["Location", "Reliability"],
                [[loc["location"], loc.get("reliability", "—")] for loc in pa["subcellular_locations"]])

        if pa.get("disease_associations"):
            doc.add_heading("Disease Involvement", level=2)
            _data_table(doc,
                ["Disease", "Source"],
                [[d["disease"], d.get("source", "HPA")] for d in pa["disease_associations"]])

        if pa.get("rna_fields"):
            doc.add_heading("RNA Expression Summary", level=2)
            _two_col_table(doc, [(r["label"], r["value"]) for r in pa["rna_fields"]])

        if pa.get("interactions"):
            doc.add_paragraph(f"Known protein interactions: {pa['interactions']}")
    else:
        doc.add_heading("12. Human Protein Atlas", level=1)
        doc.add_paragraph("No Human Protein Atlas data retrieved for this gene.")

    # ── 13. Open Targets Platform ───────────────────────────────────────
    ot = report.get("open_targets_panel")
    if ot:
        doc.add_heading("13. Open Targets Platform", level=1)
        doc.add_paragraph(
            f"Ensembl ID: {ot.get('ensembl_id','—')}  ·  "
            f"Biotype: {ot.get('biotype','—')}  ·  "
            f"{ot.get('total_disease_count',0)} total disease associations  ·  "
            f"{ot.get('total_drug_count',0)} known drugs"
        )
        if ot.get("approved_name"):
            doc.add_paragraph(f"Approved name: {ot['approved_name']}")

        if ot.get("function_descriptions"):
            doc.add_heading("Function Descriptions", level=2)
            for desc in ot["function_descriptions"]:
                doc.add_paragraph(desc, style="List Bullet")

        if ot.get("disease_rows"):
            doc.add_heading(
                f"Disease Associations (top {len(ot['disease_rows'])} of {ot['total_disease_count']})",
                level=2)
            _data_table(doc,
                ["Disease", "Association Score", "Therapeutic Areas"],
                [[r["disease_name"], str(r["score"]), r.get("therapeutic_areas","—")]
                 for r in ot["disease_rows"]])

        if ot.get("drug_rows"):
            doc.add_heading(f"Known Drugs ({ot['total_drug_count']} total)", level=2)
            _data_table(doc,
                ["Drug", "Type", "Max Phase", "Mechanism", "Indication"],
                [[r["drug_name"], r.get("drug_type","—"), r.get("phase_label","—"),
                  r.get("mechanism","—"), r.get("indication","—")]
                 for r in ot["drug_rows"]])

        if ot.get("pathways"):
            doc.add_heading("Pathway Memberships", level=2)
            _data_table(doc,
                ["Pathway", "Top-level Term"],
                [[p["name"], p.get("top_level","—")] for p in ot["pathways"]])
    else:
        doc.add_heading("13. Open Targets Platform", level=1)
        doc.add_paragraph("No Open Targets data retrieved for this gene.")

    # ── 14. Feature Annotations ─────────────────────────────────────────
    sec_n = 14
    interval_feats = report.get("interval_features") or []
    site_feats     = report.get("site_features") or []
    if interval_feats or site_feats:
        doc.add_heading(f"{sec_n}. Feature Annotations", level=1)
        sec_n += 1
        if interval_feats:
            doc.add_heading("Interval Features (Domains / Regions)", level=2)
            _data_table(doc,
                ["Source", "Type", "Label", "Coordinates", "Variant Overlap"],
                [[f["source"], f["type"], f["label"],
                  f"{f['start']}-{f['end']}", "Yes" if f.get("variant_hit") else "No"]
                 for f in interval_feats])
        if site_feats:
            doc.add_heading("Site Features (PTMs / Binding Sites)", level=2)
            _data_table(doc,
                ["Source", "Type", "Label", "Position", "Evidence"],
                [[f["source"], f["type"], f["label"],
                  str(f["position"]), f.get("evidence","—") or "—"]
                 for f in site_feats])

    # ── 15. Evidence Rows ───────────────────────────────────────────────
    evidence = report.get("evidence_rows") or []
    if evidence:
        doc.add_heading(f"{sec_n}. Evidence Annotations", level=1)
        sec_n += 1
        _data_table(doc,
            ["Source", "Label", "Position", "Summary"],
            [[e["source"], e["label"], str(e["position"]), e.get("summary","—")] for e in evidence])

    # ── 16. Supplementary Predictor Tables ─────────────────────────────
    pred_tables = report.get("supplementary_predictor_tables") or []
    if pred_tables:
        doc.add_heading(f"{sec_n}. Supplementary Predictor Tables", level=1)
        sec_n += 1
        for tbl in pred_tables:
            doc.add_heading(tbl.get("title","Table"), level=2)
            cols = tbl.get("columns", [])
            rows_t = [[str(row.get(c,"—")) for c in cols] for row in tbl.get("rows",[])]
            _data_table(doc, cols, rows_t[:50])

    # ── 17. Source Provenance ───────────────────────────────────────────
    prov = report.get("source_provenance") or []
    if prov:
        doc.add_heading(f"{sec_n}. Data Sources & Provenance", level=1)
        _data_table(doc, ["Source", "Purpose"],
                    [[p["source"], p.get("purpose","—")] for p in prov])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Chart helpers ─────────────────────────────────────────────────────────

def _metric_delta_chart(metric_deltas: list) -> bytes | None:
    if not _MPL_AVAILABLE:
        return None
    items = [(m.label, m.delta) for m in metric_deltas if m.delta is not None]
    if not items:
        return None
    labels, values = zip(*items)
    colors = ["#e05c3a" if v > 0 else "#0b6b57" for v in values]
    fig, ax = plt.subplots(figsize=(6.5, max(2.5, len(labels) * 0.45)))
    bars = ax.barh(labels, values, color=colors, height=0.6, edgecolor="white", linewidth=0.5)
    ax.axvline(0, color="#333", linewidth=0.8)
    ax.set_xlabel("Δ (mutant − wild-type)", fontsize=9)
    ax.set_title(f"Physicochemical changes: {metric_deltas[0].label[:6]}… substitution" if metric_deltas else "Metric Deltas", fontsize=10)
    ax.tick_params(labelsize=8)
    for bar, val in zip(bars, values):
        xpos = val + 0.003 if val >= 0 else val - 0.003
        ha = "left" if val >= 0 else "right"
        ax.text(xpos, bar.get_y() + bar.get_height() / 2,
                f"{val:+.4f}", va="center", ha=ha, fontsize=7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _ss_pie_chart(state_cards: list) -> bytes | None:
    if not _MPL_AVAILABLE or not state_cards:
        return None
    color_map = {"helix": "#c65d44", "strand": "#356d9c", "coil": "#7a8a78"}
    items = [(c["label"], c["count"], color_map.get(c["class_name"], "#aaa"))
             for c in state_cards if c["count"] > 0]
    if not items:
        return None
    labels, values, colors = zip(*items)
    fig, ax = plt.subplots(figsize=(3.8, 3.8))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors,
        autopct="%1.1f%%", startangle=90,
        textprops={"fontsize": 9}, pctdistance=0.75
    )
    for at in autotexts:
        at.set_fontsize(8)
    ax.set_title("Secondary structure composition", fontsize=10)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ── Document formatting helpers ───────────────────────────────────────────

def _set_page_margins(doc: "Document") -> None:
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)


def _run(paragraph, text: str, bold: bool = False, size: int | None = None,
         color: tuple | None = None, mono: bool = False) -> None:
    r = paragraph.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if mono:
        r.font.name = "Courier New"


def _two_col_table(doc: "Document", rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        c0 = table.cell(i, 0)
        c0.text = str(label)
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
        table.cell(i, 1).text = str(value) if value is not None else "—"
    doc.add_paragraph()


def _data_table(doc: "Document", columns: list[str], rows: list[list]) -> None:
    if not rows:
        doc.add_paragraph("No data available.")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(columns):
        hdr[j].text = str(col)
        for p in hdr[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = str(val) if val is not None else "—"
    doc.add_paragraph()
